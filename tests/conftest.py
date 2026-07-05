import io
from pathlib import Path

import fitz
import pytest
from docx import Document
from PIL import Image


def _make_png(width: int, height: int, color: tuple[int, int, int] = (200, 50, 50)) -> bytes:
    im = Image.new("RGB", (width, height), color)
    buf = io.BytesIO()
    im.save(buf, format="PNG")
    return buf.getvalue()


@pytest.fixture
def make_png():
    return _make_png


@pytest.fixture
def text_pdf(tmp_path) -> Path:
    doc = fitz.open()
    for i in range(2):
        page = doc.new_page(width=595, height=842)
        text = (
            f"Заголовок страницы {i + 1}\n\n"
            "Это первый абзац текста на этой странице, вполне достаточной длины, "
            "чтобы пройти порог MIN_TEXT_CHARS без проблем.\n\n"
            "Это второй абзац с дополнительным содержимым для теста извлечения текста."
        )
        page.insert_text((72, 72), text, fontsize=11)
    path = tmp_path / "text.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def scan_pdf(tmp_path) -> Path:
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    png = _make_png(595, 842)
    page.insert_image(fitz.Rect(0, 0, 595, 842), stream=png)
    path = tmp_path / "scan.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def mixed_pdf(tmp_path) -> Path:
    doc = fitz.open()
    page1 = doc.new_page(width=595, height=842)
    page1.insert_text(
        (72, 72), "Достаточно длинный текст на первой странице документа. " * 3, fontsize=11
    )
    page2 = doc.new_page(width=595, height=842)
    png = _make_png(595, 842)
    page2.insert_image(fitz.Rect(0, 0, 595, 842), stream=png)
    path = tmp_path / "mixed.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def pdf_with_inline_images(tmp_path) -> Path:
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_text((72, 72), "Текст над крупной картинкой в документе. " * 5, fontsize=11)
    big_png = _make_png(400, 400)
    page.insert_image(fitz.Rect(72, 150, 472, 550), stream=big_png)
    small_png = _make_png(10, 10)
    page.insert_image(fitz.Rect(500, 700, 510, 710), stream=small_png)
    path = tmp_path / "inline_images.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def long_text_pdf(tmp_path) -> Path:
    doc = fitz.open()
    page1 = doc.new_page(width=595, height=842)
    paragraph = (
        "Dies ist ein ausreichend langer Absatz, der wiederholt wird, damit die "
        "extrahierte Seite die Grenze fuer eine Pseudoseite deutlich ueberschreitet. "
    )
    y = 72
    for i in range(14):
        page1.insert_text((72, y), f"Absatz {i + 1}: {paragraph}", fontsize=9)
        y += 20
    png = _make_png(150, 150)
    page1.insert_image(fitz.Rect(72, y + 10, 222, 160 + y), stream=png)

    page2 = doc.new_page(width=595, height=842)
    page2.insert_text((72, 72), "Kurzer zweiter Seiteninhalt mit genug Text fuer die Erkennung.", fontsize=11)

    path = tmp_path / "long_text.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def sample_docx(tmp_path) -> Path:
    document = Document()
    document.add_paragraph("Первый абзац текста документа.")
    document.add_paragraph("Второй абзац с более длинным содержимым для проверки извлечения.")
    document.add_paragraph("Пункт списка один", style="List Bullet")
    document.add_paragraph("Пункт списка два", style="List Bullet")
    path = tmp_path / "sample.docx"
    document.save(str(path))
    return path


@pytest.fixture
def sample_image(tmp_path) -> Path:
    path = tmp_path / "sample.png"
    path.write_bytes(_make_png(100, 80))
    return path


@pytest.fixture
def sample_text_file(tmp_path) -> Path:
    path = tmp_path / "sample.txt"
    path.write_text("Первый абзац.\n\nВторой абзац.\n\nТретий абзац.", encoding="utf-8")
    return path
